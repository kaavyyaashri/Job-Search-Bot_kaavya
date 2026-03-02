# ============================================================
# Job Search Bot — Kaavya Sri Ramarapu
# Pipeline:
#   Fetch -> Local pre-filter
#   -> Gemini Call 1: Score every job against your 3 resumes
#   -> Gemini Call 2: Pick top 10 freshest best matches
#   -> Email: Top 10 highlighted + ALL filtered jobs below
# ============================================================

import google.generativeai as genai
import smtplib
import requests
import os
import sys
import time
import xml.etree.ElementTree as ET
from docx import Document
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.utils import parsedate_to_datetime
from datetime import datetime, timezone, timedelta

# ──────────────────────────────────────────────
# YOUR PROFILE (used in both Gemini calls)
# ──────────────────────────────────────────────
YOUR_PROFILE = """
Name: Kaavya Sri Ramarapu
Location: Austin, TX | OPT / STEM OPT | New Graduate (May 2025)
Degree: MS Electrical Engineering, Texas State University

Target Roles:
1. Product Engineer (hardware/software bring-up, validation, lifecycle)
2. Test & Validation Engineer (PCB testing, embedded systems, debugging)
3. Applied AI Engineer (ML/AI in semiconductor or industrial automation)
4. Graduate / Rotational Programs

Key Skills: Python, C, Linux, SLURM, PyTorch, TensorFlow, deep learning,
PCB testing, embedded systems, Arduino, FPGA, HPC clusters, root cause analysis,
cross-functional collaboration, statistical methods, computer vision

Work Authorization:
- On OPT -- authorized to work in USA now, no sponsorship needed yet
- NOT US Citizen or Green Card holder
- NOT currently enrolled (graduated May 2025)

HARD SKIP if job says:
- US Citizen only / Green Card required / Security clearance required
- Must be enrolled / Active student required

Preferred locations: Austin TX, Los Angeles CA, Boston MA, Seattle WA,
San Francisco CA, Remote USA
"""

# ──────────────────────────────────────────────
# SETTINGS
# ──────────────────────────────────────────────
ACTIVE_COUNTRIES = {
    "usa":       True,
    "india":     True,
    "singapore": True,
    "ireland":   True,
}

SEARCH_LOCATIONS = {
    "usa":       ["Austin, TX", "Los Angeles, CA", "Boston, MA", "Seattle, WA", "San Francisco, CA", "United States"],
    "india":     ["Hyderabad, India", "Bangalore, India"],
    "singapore": ["Singapore"],
    "ireland":   ["Dublin, Ireland"],
}

ROLE_KEYWORDS = [
    "product engineer new grad",
    "product engineer entry level",
    "test engineer new grad",
    "validation engineer entry level",
    "test validation engineer semiconductor",
    "applied AI engineer entry level",
    "machine learning engineer new grad",
    "rotational engineer program",
    "graduate engineer program",
    "hardware engineer new grad",
    "embedded engineer entry level",
]

# Resume file paths — these must exist in the resumes/ folder in your repo
RESUME_FILES = {
    "product":    "resumes/01_Kaavya_Sri_Resume_2026_product.docx",
    "test":       "resumes/01_Kaavya_Sri_Resume_2026_Test_Validation.docx",
    "ai":         "resumes/01_Kaavya_Sri_Resume_2026_AI_HPC.docx",
}

# ──────────────────────────────────────────────
# HOW THIS BOT WORKS
# ──────────────────────────────────────────────
# STEP 1: Fetch jobs from Indeed + Greenhouse
# STEP 2: Local keyword pre-filter (zero API, instant, cuts noise)
# STEP 3: GEMINI CALL 1 -- Resume scoring
#         Send: all pre-filtered jobs + your 3 resume texts
#         Get back: each job scored 0-100 against your actual resume content
#         + which resume to use + work auth flag
# STEP 4: GEMINI CALL 2 -- Top 10 selection
#         Send: the Gemini-scored job list
#         Get back: top 10 picks prioritizing freshness + score + OPT compatibility
# STEP 5: Send email
#         Section A: TOP 10 PICKS (highlighted, with resume match reason)
#         Section B: ALL OTHER filtered jobs (compact list with apply links)
# ──────────────────────────────────────────────


# ══════════════════════════════════════════════
# RESUME LOADER
# ══════════════════════════════════════════════

def load_resume_text(filepath):
    """Read a .docx resume and return extracted text."""
    if not os.path.exists(filepath):
        print(f"  WARNING: Resume not found: {filepath}")
        return f"[Resume file missing: {filepath}]"
    try:
        doc = Document(filepath)
        lines = []
        for para in doc.paragraphs:
            t = para.text.strip()
            if t:
                lines.append(t)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    t = cell.text.strip()
                    if t and t not in lines:
                        lines.append(t)
        text = "\n".join(lines)
        # Trim to ~1800 chars — enough context, keeps prompt size manageable
        if len(text) > 1800:
            text = text[:1800] + "\n...[trimmed]"
        print(f"  Loaded: {filepath} ({len(text)} chars)")
        return text
    except Exception as e:
        print(f"  ERROR reading {filepath}: {e}")
        return f"[Could not read resume: {e}]"


def load_all_resumes():
    print("\n--- Loading resumes ---")
    resumes = {}
    for role, path in RESUME_FILES.items():
        resumes[role] = load_resume_text(path)
    return resumes


# ══════════════════════════════════════════════
# STEP 1: FETCH JOBS
# ══════════════════════════════════════════════

def clean_text(text):
    """Strip non-ASCII to prevent encoding errors."""
    return text.encode("ascii", "ignore").decode("ascii").strip()


def is_fresh_enough(pub_date_str, max_hours=24):
    if not pub_date_str:
        return True
    try:
        pub_dt = parsedate_to_datetime(pub_date_str)
        return (datetime.now(timezone.utc) - pub_dt) <= timedelta(hours=max_hours)
    except Exception:
        return True


def fetch_from_indeed(queries):
    jobs = []
    headers = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"}
    for query in queries:
        base = "https://www.indeed.com/rss"
        loc = query["l"].lower()
        if "india" in loc:
            base = "https://in.indeed.com/rss"
        elif "singapore" in loc:
            base = "https://sg.indeed.com/rss"
        elif "ireland" in loc or "dublin" in loc:
            base = "https://ie.indeed.com/rss"

        url = f"{base}?q={requests.utils.quote(query['q'])}&l={requests.utils.quote(query['l'])}&sort=date&fromage=1"
        try:
            r = requests.get(url, headers=headers, timeout=10)
            if r.status_code != 200:
                continue
            root = ET.fromstring(r.content)
            for item in root.findall(".//item"):
                pub_date = item.findtext("pubDate", "")
                if not is_fresh_enough(pub_date):
                    continue
                jobs.append({
                    "title":    clean_text(item.findtext("title", "")),
                    "company":  clean_text(item.findtext("source", "")),
                    "link":     item.findtext("link", "").strip(),
                    "desc":     clean_text(item.findtext("description", ""))[:400],
                    "date":     pub_date,
                    "source":   "Indeed",
                    "location": query["l"],
                })
            time.sleep(0.5)
        except Exception:
            pass
    return jobs


def fetch_from_greenhouse():
    jobs = []
    companies = [
        "nvidia", "amd", "qualcomm", "applied-materials", "lam-research",
        "kla", "marvell", "lattice-semiconductor",
        "boston-dynamics", "automation-anywhere", "zebra-technologies",
        "cognex", "rockwell-automation",
        "openai", "anthropic", "scale-ai", "databricks", "datadog",
        "stripe", "hubspot", "okta", "cloudflare",
    ]
    keywords = [
        "product engineer", "test engineer", "validation engineer",
        "applied ai", "ai engineer", "machine learning", "rotational",
        "graduate", "new grad", "entry level", "hardware engineer",
        "embedded", "semiconductor", "automation engineer",
    ]
    for company in companies:
        try:
            r = requests.get(f"https://boards-api.greenhouse.io/v1/boards/{company}/jobs", timeout=8)
            if r.status_code != 200:
                continue
            for job in r.json().get("jobs", []):
                title = job.get("title", "").lower()
                if not any(kw in title for kw in keywords):
                    continue
                loc = job.get("location", {}).get("name", "") if job.get("location") else ""
                jobs.append({
                    "title":    job.get("title", ""),
                    "company":  company.replace("-", " ").title(),
                    "link":     job.get("absolute_url", ""),
                    "desc":     f"Location: {loc}.",
                    "date":     job.get("updated_at", ""),
                    "source":   "Greenhouse",
                    "location": loc,
                })
            time.sleep(0.3)
        except Exception:
            pass
    print(f"  Greenhouse: {len(jobs)} jobs")
    return jobs


def detect_country(location_str):
    loc = location_str.lower()
    if any(x in loc for x in ["india", "hyderabad", "bangalore", "bengaluru"]):
        return "india"
    elif "singapore" in loc:
        return "singapore"
    elif any(x in loc for x in ["ireland", "dublin"]):
        return "ireland"
    return "usa"


def fetch_all_jobs():
    print("\n--- STEP 1: Fetching jobs ---")
    all_jobs = []
    queries = []
    for country, active in ACTIVE_COUNTRIES.items():
        if active:
            for loc in SEARCH_LOCATIONS[country]:
                for role in ROLE_KEYWORDS:
                    queries.append({"q": role, "l": loc})

    indeed_jobs = fetch_from_indeed(queries)
    print(f"  Indeed: {len(indeed_jobs)} fresh jobs")
    all_jobs += indeed_jobs
    all_jobs += fetch_from_greenhouse()

    seen = set()
    unique = []
    for job in all_jobs:
        key = job.get("link") or job.get("title", "")
        if key and key not in seen:
            seen.add(key)
            unique.append(job)

    print(f"  Total unique: {len(unique)}")
    return unique


# ══════════════════════════════════════════════
# STEP 2: LOCAL PRE-FILTER (zero API, instant)
# ══════════════════════════════════════════════

def local_prefilter(all_jobs, top_n_per_country=20):
    """
    Fast local keyword scoring — no API call.
    Removes obvious mismatches (senior roles, citizenship required)
    and keeps top N per country to feed into Gemini Call 1.
    """
    GOOD = [
        "new grad", "entry level", "graduate", "rotational", "early career",
        "associate engineer", "product engineer", "test engineer",
        "validation engineer", "applied ai", "ai engineer",
        "machine learning", "hardware engineer", "embedded",
        "python", "pytorch", "tensorflow", "linux", "hpc", "slurm",
        "pcb", "fpga", "semiconductor", "automation", "deep learning",
        "nvidia", "amd", "qualcomm", "texas instruments", "applied materials",
        "lam research", "kla", "siemens", "honeywell", "rockwell", "cognex",
    ]
    SKIP = [
        "senior", "sr.", "staff engineer", "principal engineer",
        "director", "manager", "vp ", "vice president",
        "5+ years", "7+ years", "8+ years", "10+ years",
        "us citizen only", "clearance required", "secret clearance",
        "must be enrolled", "currently pursuing",
    ]

    buckets = {"usa": [], "india": [], "singapore": [], "ireland": []}

    for job in all_jobs:
        combined = (job.get("title","") + " " + job.get("desc","") + " " + job.get("company","")).lower()
        if any(kw in combined for kw in SKIP):
            continue
        score = sum(2 if kw in job.get("title","").lower() else 1
                    for kw in GOOD if kw in combined)
        if score == 0:
            continue
        country = detect_country(job.get("location",""))
        if ACTIVE_COUNTRIES.get(country, False):
            buckets[country].append((score, job))

    result = []
    for country, scored in buckets.items():
        scored.sort(key=lambda x: x[0], reverse=True)
        for score, job in scored[:top_n_per_country]:
            job["country"] = country
            job["local_score"] = score
            result.append(job)

    print(f"\n--- STEP 2: Local pre-filter kept {len(result)} jobs ---")
    for c in ["usa", "india", "singapore", "ireland"]:
        n = sum(1 for j in result if j.get("country") == c)
        if n:
            print(f"  {c.upper()}: {n} jobs")
    return result


# ══════════════════════════════════════════════
# GEMINI CALL 1 — Resume keyword matching & scoring
# ══════════════════════════════════════════════

def gemini_call_1_score_jobs(prefiltered_jobs, resumes):
    """
    GEMINI CALL 1:
    Sends every pre-filtered job + all 3 resume texts to Gemini.
    Gemini reads your actual resume content and scores each job 0-100
    based on how well your real experience matches the job description.
    Also flags work authorization and which resume to use.

    Returns: list of job dicts with added 'resume_score', 'resume_to_use', 'work_auth'
    """
    if not prefiltered_jobs:
        return []

    genai.configure(api_key=os.environ["GEMINI_API_KEY"])
    model = genai.GenerativeModel("gemini-2.0-flash-lite")

    # Build compact job list
    jobs_text = ""
    for i, job in enumerate(prefiltered_jobs, 1):
        jobs_text += (
            f"\nJOB {i}: [{job.get('country','').upper()}] {job['title']} "
            f"at {job.get('company','?')} | {job.get('location','')}\n"
            f"Desc: {job.get('desc','')[:250]}\n"
            f"Link: {job['link']}\n"
        )

    # Include all 3 resume summaries
    resume_block = f"""
=== RESUME 1: Product Engineer ===
{resumes.get('product', '[missing]')}

=== RESUME 2: Test & Validation Engineer ===
{resumes.get('test', '[missing]')}

=== RESUME 3: Applied AI / HPC Engineer ===
{resumes.get('ai', '[missing]')}
"""

    prompt = f"""You are a job matching expert. Score each job against the candidate's actual resume experience.

CANDIDATE PROFILE:
{YOUR_PROFILE}

CANDIDATE'S RESUMES (3 versions, each targeting a different role type):
{resume_block}

JOB LISTINGS TO SCORE ({len(prefiltered_jobs)} jobs):
{jobs_text}

YOUR TASK:
For each job:
1. Read the job description and the 3 resumes carefully
2. Score the job 0-100 based on how well the candidate's ACTUAL experience matches
   (40pts role fit, 30pts skills match from resume, 20pts company type, 10pts location)
3. Pick which resume fits best: "product", "test", or "ai"
4. Flag work authorization: "OPT_OK" / "VERIFY" / "SKIP"
   - OPT_OK: no citizenship requirement mentioned, or says visa sponsorship available
   - VERIFY: says no sponsorship (may still accept OPT -- candidate should confirm)
   - SKIP: explicitly says US Citizen only or Green Card required or clearance required
5. Write one short sentence on why the score is what it is

RESPOND IN THIS EXACT FORMAT (one line per job, pipe-delimited, no extra text):
JOB_NUM|SCORE|RESUME|WORK_AUTH|REASON

Example:
1|78|product|OPT_OK|Strong PCB testing and Python automation background matches this validation role.
2|45|ai|VERIFY|ML skills fit but no semiconductor industry experience yet.
3|15|test|SKIP|Requires US citizenship and 5 years experience.

Output exactly {len(prefiltered_jobs)} lines. No headers. No markdown. Just the lines.
"""

    print(f"\n--- GEMINI CALL 1: Scoring {len(prefiltered_jobs)} jobs against your resumes ---")
    raw = _call_gemini(model, prompt, call_name="Call 1 (Resume Scoring)")
    if not raw:
        print("  Gemini Call 1 failed -- using local scores as fallback")
        for job in prefiltered_jobs:
            job["resume_score"] = job.get("local_score", 0) * 5
            job["resume_to_use"] = "product"
            job["work_auth"] = "VERIFY"
            job["score_reason"] = "Matched resume keywords locally (Gemini unavailable)."
        return prefiltered_jobs

    # Parse the response and attach scores back to job dicts
    score_map = {}
    for line in raw.strip().split("\n"):
        line = line.strip()
        if not line or "|" not in line:
            continue
        parts = line.split("|")
        if len(parts) < 5:
            continue
        try:
            job_num  = int(parts[0].strip()) - 1  # 0-indexed
            score    = int(parts[1].strip())
            resume   = parts[2].strip().lower()
            work_auth = parts[3].strip()
            reason   = parts[4].strip()
            score_map[job_num] = (score, resume, work_auth, reason)
        except Exception:
            continue

    # Attach scores to jobs, skip SKIP jobs
    scored_jobs = []
    for i, job in enumerate(prefiltered_jobs):
        if i in score_map:
            score, resume, work_auth, reason = score_map[i]
            if work_auth == "SKIP":
                continue  # Remove citizenship-required jobs
            job["resume_score"] = score
            job["resume_to_use"] = resume
            job["work_auth"] = work_auth
            job["score_reason"] = reason
        else:
            job["resume_score"] = job.get("local_score", 0) * 5
            job["resume_to_use"] = "product"
            job["work_auth"] = "VERIFY"
            job["score_reason"] = "Score estimated from local keywords."
        scored_jobs.append(job)

    # Sort all by resume_score descending
    scored_jobs.sort(key=lambda x: x.get("resume_score", 0), reverse=True)

    print(f"  Gemini scored {len(scored_jobs)} jobs (removed citizenship-required)")
    return scored_jobs


# ══════════════════════════════════════════════
# GEMINI CALL 2 — Pick top 10 fresh jobs
# ══════════════════════════════════════════════

def gemini_call_2_pick_top10(scored_jobs):
    """
    GEMINI CALL 2:
    Sends the already-scored job list to Gemini.
    Asks Gemini to pick the TOP 10 considering:
    - Resume score (higher = better match)
    - Freshness (posted today = priority)
    - OPT_OK jobs get priority over VERIFY
    - Geographic spread (not all from same location)

    Returns: list of top 10 job dicts with added 'rank' and 'pick_reason'
    """
    if not scored_jobs:
        return []

    genai.configure(api_key=os.environ["GEMINI_API_KEY"])
    model = genai.GenerativeModel("gemini-2.0-flash-lite")

    # Build scored job list for prompt
    jobs_text = ""
    for i, job in enumerate(scored_jobs, 1):
        jobs_text += (
            f"\n{i}. Score:{job.get('resume_score',0)} | {job.get('work_auth','?')} | "
            f"[{job.get('country','').upper()}] {job['title']} at {job.get('company','?')} "
            f"| {job.get('location','')} | Posted: {job.get('date','unknown')}\n"
            f"   Resume match reason: {job.get('score_reason','')}\n"
            f"   Link: {job['link']}\n"
        )

    prompt = f"""You already scored these jobs against the candidate's resume. Now pick the final TOP 10.

CANDIDATE: Kaavya Sri Ramarapu | OPT | MS EE | New Grad May 2025
Austin TX preferred | Product Eng / Test Eng / Applied AI

SCORED JOBS (sorted by resume score, highest first):
{jobs_text}

YOUR TASK:
Pick the TOP 10 jobs to recommend. Prioritize in this order:
1. Jobs posted TODAY or very recently (freshness matters most -- apply before competition)
2. Higher resume score (better match to actual experience)
3. OPT_OK over VERIFY
4. Geographic spread -- include USA, India, Singapore, Ireland if available
5. Mix of role types -- not all product engineer, mix in test and AI roles

For each pick, write ONE sentence on why you chose it (combine freshness + score reason).

RESPOND IN THIS EXACT FORMAT (pipe-delimited, no headers, no markdown, exactly 10 lines):
RANK|ORIGINAL_NUM|COUNTRY|TITLE|COMPANY|LOCATION|RESUME|WORK_AUTH|WHY_CHOSEN|LINK

Example:
1|3|USA|Product Engineer|NVIDIA|Austin TX|product|OPT_OK|Posted today with strong PCB and Python match from resume.|https://...
2|7|INDIA|Test Engineer|Qualcomm|Hyderabad|test|OPT_OK|Fresh listing, embedded validation skills align directly.|https://...
"""

    print(f"\n--- GEMINI CALL 2: Picking top 10 from {len(scored_jobs)} scored jobs ---")
    raw = _call_gemini(model, prompt, call_name="Call 2 (Top 10 Selection)")
    if not raw:
        print("  Gemini Call 2 failed -- using top 10 by score as fallback")
        top10 = []
        for i, job in enumerate(scored_jobs[:10], 1):
            job["rank"] = i
            job["pick_reason"] = job.get("score_reason", "High resume match score.")
            top10.append(job)
        return top10

    # Parse top 10 response
    top10 = []
    for line in raw.strip().split("\n"):
        line = line.strip()
        if not line or "|" not in line:
            continue
        parts = line.split("|")
        if len(parts) < 10:
            continue
        try:
            orig_idx = int(parts[1].strip()) - 1
            job = scored_jobs[orig_idx].copy() if orig_idx < len(scored_jobs) else {}
            job["rank"]        = parts[0].strip()
            job["country"]     = parts[2].strip()
            job["title"]       = parts[3].strip()
            job["company"]     = parts[4].strip()
            job["location"]    = parts[5].strip()
            job["resume_to_use"] = parts[6].strip()
            job["work_auth"]   = parts[7].strip()
            job["pick_reason"] = parts[8].strip()
            job["link"]        = parts[9].strip()
            top10.append(job)
        except Exception:
            continue

    print(f"  Gemini selected {len(top10)} top jobs")
    return top10


# ══════════════════════════════════════════════
# SHARED GEMINI HELPER (handles 429 retry)
# ══════════════════════════════════════════════

def _call_gemini(model, prompt, call_name=""):
    """Call Gemini with automatic retry on 429 rate limit."""
    try:
        response = model.generate_content(prompt)
        print(f"  {call_name}: OK")
        return response.text.strip()
    except Exception as e:
        err = str(e)
        if "429" in err:
            import re
            m = re.search(r'seconds:\s*(\d+)', err)
            wait = int(m.group(1)) + 10 if m else 70
            print(f"  {call_name}: Rate limited -- waiting {wait}s then retrying...")
            time.sleep(wait)
            try:
                response = model.generate_content(prompt)
                print(f"  {call_name}: Retry OK")
                return response.text.strip()
            except Exception as e2:
                print(f"  {call_name}: Retry failed: {e2}")
                return None
        print(f"  {call_name}: Error: {e}")
        return None


# ══════════════════════════════════════════════
# STEP 5: SEND EMAIL
# ══════════════════════════════════════════════

def build_resume_badge(resume_key):
    badges = {
        "product": ("Product Engineer", "#1a6e4a"),
        "test":    ("Test & Validation", "#4a1a6e"),
        "ai":      ("Applied AI / HPC",  "#6e4a1a"),
    }
    label, color = badges.get(resume_key, ("General", "#555"))
    return f'<span style="background:{color}; color:white; padding:2px 8px; border-radius:10px; font-size:11px;">{label}</span>'


def build_auth_badge(work_auth):
    if "OK" in work_auth:
        return '<span style="color:#27ae60; font-weight:bold;">OPT OK</span>'
    elif "VERIFY" in work_auth:
        return '<span style="color:#e67e22; font-weight:bold;">VERIFY</span>'
    else:
        return '<span style="color:#e74c3c; font-weight:bold;">SKIP</span>'


def send_email(top10_jobs, all_scored_jobs, total_fetched):
    sender_email    = os.environ["EMAIL_ADDRESS"]
    sender_password = os.environ["EMAIL_APP_PASSWORD"]
    recipient_email = os.environ["EMAIL_TO"]

    today    = datetime.now().strftime("%B %d, %Y")
    run_time = datetime.now().strftime("%Y-%m-%d %H:%M UTC")

    msg = MIMEMultipart("alternative")
    msg["Subject"] = f"Job Digest | Top {len(top10_jobs)} Picks + {len(all_scored_jobs)} Total | {today}"
    msg["From"]    = sender_email
    msg["To"]      = recipient_email

    # ── Section A: TOP 10 PICKS ──────────────────────────────
    top10_rows = ""
    for job in top10_jobs:
        resume_badge = build_resume_badge(job.get("resume_to_use", "product"))
        auth_badge   = build_auth_badge(job.get("work_auth", "VERIFY"))
        score        = job.get("resume_score", 0)
        score_color  = "#27ae60" if score >= 70 else ("#e67e22" if score >= 45 else "#e74c3c")

        top10_rows += f"""
        <tr style="border-bottom:1px solid #eee;">
          <td style="padding:14px 10px; font-size:16px; font-weight:bold;
                     color:#4a4a8a; width:35px;">#{job.get('rank','?')}</td>
          <td style="padding:14px 10px;">
            <div style="font-size:14px; font-weight:bold; color:#1a1a2e;">
              {job.get('title','')}
            </div>
            <div style="font-size:12px; color:#666; margin-top:3px;">
              {job.get('company','')} &middot; {job.get('location','')}
            </div>
            <div style="margin-top:5px;">
              {resume_badge}
            </div>
          </td>
          <td style="padding:14px 10px; font-size:12px; color:#444;">
            {job.get('pick_reason','')}
          </td>
          <td style="padding:14px 10px; text-align:center;">
            <div style="font-size:16px; font-weight:bold; color:{score_color};">
              {score}
            </div>
            <div style="font-size:10px; color:#999;">/100</div>
          </td>
          <td style="padding:14px 10px; text-align:center;">
            {auth_badge}
          </td>
          <td style="padding:14px 10px;">
            <a href="{job.get('link','')}"
               style="background:#1a1a2e; color:white; padding:8px 16px;
                      border-radius:6px; text-decoration:none; font-size:12px;
                      white-space:nowrap; display:inline-block;">
              Apply Now
            </a>
          </td>
        </tr>"""

    # ── Section B: ALL OTHER FILTERED JOBS ───────────────────
    # Jobs not in top 10, grouped by country
    top10_links = {j.get("link","") for j in top10_jobs}
    remaining = [j for j in all_scored_jobs if j.get("link","") not in top10_links]

    country_labels = {"usa": "United States", "india": "India",
                      "singapore": "Singapore", "ireland": "Ireland"}
    country_flags  = {"usa": "US", "india": "IN", "singapore": "SG", "ireland": "IE"}

    all_jobs_sections = ""
    for country in ["usa", "india", "singapore", "ireland"]:
        country_jobs = [j for j in remaining if j.get("country","") == country]
        if not country_jobs:
            continue

        rows = ""
        for job in country_jobs:
            auth_badge = build_auth_badge(job.get("work_auth","VERIFY"))
            score = job.get("resume_score", 0)
            score_color = "#27ae60" if score >= 70 else ("#e67e22" if score >= 45 else "#aaa")
            rows += f"""
            <tr style="border-bottom:1px solid #f5f5f5;">
              <td style="padding:10px 8px; font-size:13px;">
                <strong>{job.get('title','')}</strong><br>
                <span style="color:#888; font-size:11px;">
                  {job.get('company','')} &middot; {job.get('location','')}
                </span>
              </td>
              <td style="padding:10px 8px; font-size:11px; color:#444;">
                {job.get('score_reason','')[:100]}
              </td>
              <td style="padding:10px 8px; font-size:13px; font-weight:bold;
                         color:{score_color}; text-align:center; width:40px;">
                {score}
              </td>
              <td style="padding:10px 8px; text-align:center; width:70px;">
                {auth_badge}
              </td>
              <td style="padding:10px 8px; width:60px;">
                <a href="{job.get('link','')}"
                   style="background:#4a4a8a; color:white; padding:5px 10px;
                          border-radius:4px; text-decoration:none; font-size:11px;">
                  Apply
                </a>
              </td>
            </tr>"""

        flag = country_flags.get(country, "")
        all_jobs_sections += f"""
        <div style="margin-bottom:24px;">
          <div style="background:#4a4a8a; color:white; padding:10px 16px;
                      border-radius:6px 6px 0 0; font-size:13px; font-weight:bold;">
            {flag} {country_labels.get(country, country.upper())}
            <span style="opacity:0.7; font-weight:normal; font-size:12px;">
              &nbsp;({len(country_jobs)} jobs)
            </span>
          </div>
          <div style="border:1px solid #ddd; border-top:none; border-radius:0 0 6px 6px; overflow:hidden;">
            <table style="width:100%; border-collapse:collapse;">
              <thead>
                <tr style="background:#f8f8f8;">
                  <th style="padding:8px; text-align:left; font-size:11px; color:#999;">Role &amp; Company</th>
                  <th style="padding:8px; text-align:left; font-size:11px; color:#999;">Why It Matches</th>
                  <th style="padding:8px; text-align:center; font-size:11px; color:#999;">Score</th>
                  <th style="padding:8px; text-align:center; font-size:11px; color:#999;">Auth</th>
                  <th style="padding:8px; font-size:11px; color:#999;">Apply</th>
                </tr>
              </thead>
              <tbody>{rows}</tbody>
            </table>
          </div>
        </div>"""

    # ── Assemble full HTML ────────────────────────────────────
    full_html = f"""<!DOCTYPE html>
<html>
<body style="font-family:Arial,sans-serif; max-width:860px; margin:auto; padding:20px; background:#f5f5f5;">

  <!-- Header -->
  <div style="background:linear-gradient(135deg,#1a1a2e,#4a4a8a); color:white;
              padding:24px; border-radius:10px; margin-bottom:20px;">
    <h1 style="margin:0 0 6px 0; font-size:20px;">Daily Engineering Job Digest</h1>
    <p style="margin:0; opacity:0.85; font-size:13px;">
      Kaavya Sri &bull; Product Eng &bull; Test &amp; Validation &bull; Applied AI
      &bull; OPT Friendly &bull; {today}
    </p>
  </div>

  <!-- Pipeline stats -->
  <div style="background:white; border-radius:8px; padding:16px 20px;
              margin-bottom:20px; font-size:13px; color:#333;
              border-left:4px solid #4a4a8a;">
    <strong>Today's pipeline:</strong>&nbsp;
    <strong>{total_fetched}</strong> total listings scraped
    &rarr; <strong>{len(all_scored_jobs)}</strong> passed resume scoring
    &rarr; <strong>{len(top10_jobs)}</strong> top picks selected by Gemini
    <br><br>
    <span style="font-size:12px; color:#666;">
      Gemini used 2 API calls: (1) scored every job against your 3 resumes,
      (2) selected top 10 prioritizing freshness + score + OPT compatibility.
    </span>
  </div>

  <!-- Work auth key -->
  <div style="background:#fff8e1; border:1px solid #ffc107; border-radius:8px;
              padding:12px 16px; margin-bottom:24px; font-size:12px; color:#555;">
    <strong>Work Auth Key:</strong>&nbsp;
    <span style="color:#27ae60; font-weight:bold;">OPT OK</span> = apply now &nbsp;|&nbsp;
    <span style="color:#e67e22; font-weight:bold;">VERIFY</span> = says no sponsorship but OPT often accepted -- confirm with recruiter &nbsp;|&nbsp;
    <span style="color:#e74c3c; font-weight:bold;">SKIP</span> = citizenship required (removed from top 10)
    <br>Check H1B history:
    <a href="https://www.myvisajobs.com" style="color:#4a4a8a;">myvisajobs.com</a>
  </div>

  <!-- SECTION A: TOP 10 PICKS -->
  <div style="background:white; border-radius:10px; overflow:hidden;
              box-shadow:0 2px 12px rgba(0,0,0,0.1); margin-bottom:32px;">
    <div style="background:#1a1a2e; color:white; padding:14px 20px;">
      <h2 style="margin:0; font-size:16px;">
        Top {len(top10_jobs)} Picks -- Apply These First
      </h2>
      <p style="margin:4px 0 0 0; opacity:0.7; font-size:12px;">
        Ranked by freshness + resume match score + OPT compatibility
      </p>
    </div>
    <table style="width:100%; border-collapse:collapse;">
      <thead>
        <tr style="background:#f8f8f8; border-bottom:2px solid #e0e0e0;">
          <th style="padding:10px 10px; text-align:left; font-size:11px; color:#999; width:35px;">#</th>
          <th style="padding:10px 10px; text-align:left; font-size:11px; color:#999;">Role &amp; Company</th>
          <th style="padding:10px 10px; text-align:left; font-size:11px; color:#999;">Why Gemini Picked It</th>
          <th style="padding:10px 10px; text-align:center; font-size:11px; color:#999; width:50px;">Score</th>
          <th style="padding:10px 10px; text-align:center; font-size:11px; color:#999; width:70px;">Work Auth</th>
          <th style="padding:10px 10px; text-align:left; font-size:11px; color:#999; width:90px;">Apply</th>
        </tr>
      </thead>
      <tbody>{top10_rows}</tbody>
    </table>
  </div>

  <!-- Before you apply checklist -->
  <div style="background:#1a1a2e; color:white; border-radius:8px;
              padding:18px 20px; margin-bottom:32px; font-size:13px;">
    <strong>Before You Apply to Any Top Pick:</strong>
    <ol style="margin:8px 0 0 0; padding-left:20px; line-height:2.2;">
      <li>Use the resume shown in the badge (Product / Test &amp; Validation / Applied AI)</li>
      <li>VERIFY jobs -- "no sponsorship" usually means no H1B transfer, OPT still often accepted</li>
      <li>Check company H1B history at myvisajobs.com</li>
      <li>Apply same day -- fresh jobs fill fast, early applicants stand out</li>
    </ol>
  </div>

  <!-- SECTION B: ALL OTHER FILTERED JOBS -->
  <div style="border-top:2px solid #ddd; padding-top:24px; margin-bottom:12px;">
    <h2 style="color:#1a1a2e; font-size:16px; margin:0 0 4px 0;">
      All Filtered Jobs ({len(remaining)} more across all regions)
    </h2>
    <p style="color:#888; font-size:12px; margin:0 0 20px 0;">
      These passed resume scoring but were not in the top 10.
      Worth checking if the top picks are not a fit.
    </p>
    {all_jobs_sections}
  </div>

  <!-- Footer -->
  <div style="text-align:center; color:#aaa; font-size:11px;
              margin-top:16px; padding-top:16px; border-top:1px solid #ddd;">
    Generated by Gemini AI (2 calls) + GitHub Actions &bull; {run_time}<br>
    Regions: USA &bull; India &bull; Singapore &bull; Ireland
  </div>

</body>
</html>"""

    plain = (
        f"Daily Job Digest - {today}\n\n"
        f"Pipeline: {total_fetched} scraped -> {len(all_scored_jobs)} scored -> {len(top10_jobs)} top picks\n\n"
        f"Open in an HTML email client to view the full formatted digest.\n\n"
        f"Run time: {run_time}"
    )

    msg.attach(MIMEText(plain, "plain", "utf-8"))
    msg.attach(MIMEText(full_html, "html", "utf-8"))

    with smtplib.SMTP_SSL("smtp.gmail.com", 465) as server:
        server.login(sender_email, sender_password)
        server.send_message(msg)   # handles Unicode/emoji -- no ASCII crash

    print(f"  Email sent to {recipient_email}")


# ══════════════════════════════════════════════
# ERROR / NO-MATCH EMAILS
# ══════════════════════════════════════════════

def send_error_email(subject, detail):
    try:
        s = os.environ.get("EMAIL_ADDRESS","")
        p = os.environ.get("EMAIL_APP_PASSWORD","")
        r = os.environ.get("EMAIL_TO","")
        if not all([s,p,r]):
            return
        msg = MIMEMultipart("alternative")
        msg["Subject"] = f"Job Bot Error: {subject}"
        msg["From"] = s
        msg["To"]   = r
        html = f"""<html><body style="font-family:Arial; max-width:600px; margin:auto; padding:20px;">
  <div style="background:#e74c3c;color:white;padding:20px;border-radius:8px;margin-bottom:16px;">
    <h2 style="margin:0;">Job Bot Error: {subject}</h2></div>
  <pre style="background:#f9f9f9;padding:16px;border-radius:6px;font-size:12px;
       white-space:pre-wrap;overflow-wrap:break-word;">{detail}</pre>
  <p style="font-size:13px;color:#666;">Debug: GitHub &rarr; Actions &rarr; failed run &rarr; expand logs.</p>
</body></html>"""
        msg.attach(MIMEText(html, "html", "utf-8"))
        with smtplib.SMTP_SSL("smtp.gmail.com", 465) as server:
            server.login(s, p)
            server.send_message(msg)
    except Exception as e:
        print(f"  Could not send error email: {e}")


def send_no_matches_email():
    try:
        s = os.environ.get("EMAIL_ADDRESS","")
        p = os.environ.get("EMAIL_APP_PASSWORD","")
        r = os.environ.get("EMAIL_TO","")
        today = datetime.now().strftime("%B %d, %Y")
        msg = MIMEMultipart("alternative")
        msg["Subject"] = f"Job Bot - No New Matches | {today}"
        msg["From"] = s
        msg["To"]   = r
        html = f"""<html><body style="font-family:Arial; max-width:600px; margin:auto; padding:20px;">
  <div style="background:#f0f0ff;border:1px solid #c0c0ff;border-radius:10px;padding:24px;text-align:center;">
    <h2 style="color:#4a4a8a;margin:0 0 8px 0;">No New Matches Today</h2>
    <p style="color:#666;margin:0;">{today}</p>
  </div>
  <p style="font-size:14px;color:#444;margin-top:16px;line-height:1.7;">
    The bot ran across all regions but no fresh listings matched your profile today.
    Normal on quiet days -- next run is scheduled soon.
  </p>
</body></html>"""
        msg.attach(MIMEText(html, "html", "utf-8"))
        with smtplib.SMTP_SSL("smtp.gmail.com", 465) as server:
            server.login(s, p)
            server.send_message(msg)
    except Exception as e:
        print(f"  Could not send no-matches email: {e}")


# ══════════════════════════════════════════════
# MAIN PIPELINE
# ══════════════════════════════════════════════

def run_pipeline():
    start = datetime.now()
    print("\n" + "="*55)
    print("  Engineering Job Search Bot - Kaavya Sri")
    print(f"  {start.strftime('%A, %B %d, %Y - %H:%M UTC')}")
    print("="*55)

    active = [c.upper() for c, v in ACTIVE_COUNTRIES.items() if v]
    print(f"\nActive regions: {', '.join(active)}")

    # Check secrets
    print("\nChecking secrets...")
    required = ["GEMINI_API_KEY", "EMAIL_ADDRESS", "EMAIL_APP_PASSWORD", "EMAIL_TO"]
    missing  = [s for s in required if not os.environ.get(s)]
    for s in required:
        print(f"  {'OK' if s not in missing else 'MISSING'} {s}")
    if missing:
        print(f"\nFATAL: Missing secrets: {', '.join(missing)}")
        sys.exit(1)

    # Load resumes
    resumes = load_all_resumes()

    # STEP 1: Fetch
    try:
        all_jobs = fetch_all_jobs()
    except Exception as e:
        print(f"Fatal fetch error: {e}")
        send_error_email("Job Fetching Failed", str(e))
        sys.exit(1)

    if not all_jobs:
        send_error_email("No Jobs Found", "No listings fetched from any source.")
        sys.exit(0)

    # STEP 2: Local pre-filter
    print("\n--- STEP 2: Local pre-filter ---")
    prefiltered = local_prefilter(all_jobs, top_n_per_country=20)
    if not prefiltered:
        send_no_matches_email()
        sys.exit(0)

    # STEP 3: Gemini Call 1 -- score every job against resumes
    # Wait 15s between calls to stay within free tier rate limits
    scored_jobs = gemini_call_1_score_jobs(prefiltered, resumes)
    if not scored_jobs:
        send_no_matches_email()
        sys.exit(0)

    print(f"  Waiting 20s before Call 2 (free tier rate limit buffer)...")
    time.sleep(20)

    # STEP 4: Gemini Call 2 -- pick top 10
    top10_jobs = gemini_call_2_pick_top10(scored_jobs)
    if not top10_jobs:
        top10_jobs = [dict(j, rank=i+1, pick_reason=j.get("score_reason",""))
                      for i, j in enumerate(scored_jobs[:10])]

    # STEP 5: Send email
    print("\n--- STEP 5: Sending email ---")
    try:
        send_email(top10_jobs, scored_jobs, len(all_jobs))
    except Exception as e:
        print(f"Email error: {e}")
        send_error_email("Email Send Failed", str(e))
        sys.exit(1)

    elapsed = (datetime.now() - start).seconds
    print("\n" + "="*55)
    print(f"  Done!")
    print(f"  {len(all_jobs)} fetched -> {len(prefiltered)} pre-filtered -> "
          f"{len(scored_jobs)} resume-scored -> {len(top10_jobs)} top picks emailed")
    print(f"  Gemini calls used: 2")
    print(f"  Runtime: {elapsed}s")
    print(f"  Email sent to: {os.environ.get('EMAIL_TO')}")
    print("="*55 + "\n")


if __name__ == "__main__":
    run_pipeline()
