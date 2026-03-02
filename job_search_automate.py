# Job Search automate code with the help of Gemini AI

import google.generativeai as genai
import smtplib
import requests
import os
from docx import Document
import xml.etree.ElementTree as ET
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from datetime import datetime
import sys
# ──────────────────────────────────────────────
# YOUR PROFILE — Edit this section
# ──────────────────────────────────────────────
YOUR_PROFILE = """
## Personal Details:
- Name: Kaavya Sri Ramarapu
- Location: Austin, TX
- Email: ramarapukaavyasri@gmail.com

## Education:
- Master of Science in Electrical Engineering — Texas State University (Aug 2022 – May 2025)
- Research Focus: Environment Emotion Classification to assist Children with ASD
- Courses: Computer Architecture, Machine Learning for Engineering Applications,
  Digital Image Processing, Engineering Economic Analysis, Statistical Methods
- Bachelor of Technology in Electrical and Electronics Engineering
  Jawaharlal Nehru Technological University, Hyderabad (Aug 2016 – Sept 2020)

## Target Roles (apply to ALL three):
1. Product Engineer — hardware/software product bring-up, validation, lifecycle improvement
2. Test & Validation Engineer — PCB testing, embedded systems validation, structured debugging
3. Applied AI Engineer — ML/AI in semiconductor or industrial automation environments

## Also Targeting:
- Graduate Programs (rotational or structured graduate hire programs)
- Rotational Engineering Programs at companies

## Key Skills:
- Languages: Python, C, Linux scripting, SLURM
- ML/AI: PyTorch, TensorFlow, deep learning, GPU-accelerated training, HPC clusters
- Hardware: PCB testing, embedded systems, hardware-software validation
- Tools: Root cause analysis, performance optimization, cross-functional collaboration

## Work Authorization:
- OPT (Optional Practical Training) — authorized to work in USA
- STEM OPT eligible
- Does NOT need H1B sponsorship right now
- NOT a US Citizen or Green Card holder
- NOT currently enrolled as a student (graduated May 2025)
- Future sponsorship is flexible and company-dependent

## Experience Level:
- New Graduate / Entry Level / 0–2 years

## Preferred Company Types:
- Startups and mid-size companies
- Semiconductor companies (NVIDIA, AMD, Qualcomm, Texas Instruments, Applied Materials,
  Lam Research, KLA, Marvell, Lattice, Microchip, Infineon, NXP, STMicroelectronics)
- Industrial Automation (Siemens, ABB, Rockwell Automation, Honeywell, Emerson,
  Cognex, National Instruments / NI, Fanuc, KUKA, Zebra Technologies)
- Companies with structured new grad / rotational programs

## Preferred Locations (in order):
1. Austin, TX
2. Los Angeles, CA
3. Boston, MA
4. Seattle, WA
5. San Francisco, CA
6. Remote (USA-based)

## HARD FILTERS — Skip any job that says:
- "US Citizen only" or "US Citizen required"
- "Green Card required" or "Permanent Resident required"
- "Security clearance required"
- "Currently pursuing a degree" or "Must be enrolled"
- "Active student enrollment required"
- "No international students"
"""

# ──────────────────────────────────────────────
# LOCATIONS — Edit this to change search cities
# ──────────────────────────────────────────────
SEARCH_LOCATIONS = {
    "usa": [
        "Austin, TX",
        "Los Angeles, CA",
        "Boston, MA",
        "Seattle, WA",
        "San Francisco, CA",
        "United States",        # catches remote + nationwide posts
    ],
    "india": [
        "Hyderabad, India",
        "Bangalore, India",
    ],
    "singapore": [
        "Singapore",
    ],
    "ireland": [
        "Dublin, Ireland",
    ],
}

# ── Toggle countries ON/OFF here ──────────────
# Set to True to search, False to skip
ACTIVE_COUNTRIES = {
    "usa":       True,
    "india":     True,
    "singapore": True,
    "ireland":   True,
}

# ──────────────────────────────────────────────
# ROLE KEYWORDS — Edit to add/remove role searches
# ──────────────────────────────────────────────
ROLE_KEYWORDS = [
    "product engineer new grad",
    "product engineer entry level",
    "product engineer rotational program",
    "test engineer new grad",
    "validation engineer entry level",
    "test validation engineer semiconductor",
    "test validation engineer industrial automation",
    "applied AI engineer semiconductor",
    "AI engineer industrial automation entry level",
    "machine learning engineer semiconductor new grad",
    "rotational engineer program new grad",
    "graduate engineer program",
    "engineering rotational program",
]

# ──────────────────────────────────────────────
# AUTO-BUILD SEARCH QUERIES from above settings
# ──────────────────────────────────────────────
SEARCH_QUERIES = []

for country, is_active in ACTIVE_COUNTRIES.items():
    if is_active:
        for location in SEARCH_LOCATIONS[country]:
            for role in ROLE_KEYWORDS:
                SEARCH_QUERIES.append({"q": role, "l": location})

# ──────────────────────────────────────────────
# COMPANIES KNOWN TO SPONSOR — Used as a bonus filter
# ──────────────────────────────────────────────
KNOWN_SPONSORS = [
    "Google", "Microsoft", "Amazon", "Meta", "Apple", "Salesforce",
    "Stripe", "Airbnb", "Lyft", "Uber", "DoorDash", "Coinbase",
    "Databricks", "Snowflake", "Figma", "Notion", "Linear", "Vercel",
    "OpenAI", "Anthropic", "Scale AI", "Palantir", "Twilio", "MongoDB",
    "Atlassian", "Asana", "HubSpot", "Zendesk", "Okta", "Cloudflare",
    "Datadog", "HashiCorp", "Splunk", "New Relic", "PagerDuty",
]

# ──────────────────────────────────────────────
# STEP 1: Fetch jobs for ALL countries
# ──────────────────────────────────────────────
import time
from datetime import datetime, timezone, timedelta
from email.utils import parsedate_to_datetime

# ──────────────────────────────────────────────
# COUNTRY-SPECIFIC JOB BOARD RSS URLS
# ──────────────────────────────────────────────

def build_indeed_urls(queries):
    """Build Indeed RSS URLs — works for USA, India, Singapore, Ireland"""
    urls = []
    for query in queries:
        base = "https://www.indeed.com/rss"
        # Use country-specific Indeed domains
        location = query["l"].lower()
        if "india" in location:
            base = "https://in.indeed.com/rss"
        elif "singapore" in location:
            base = "https://sg.indeed.com/rss"
        elif "ireland" in location or "dublin" in location:
            base = "https://ie.indeed.com/rss"

        url = (
            f"{base}"
            f"?q={requests.utils.quote(query['q'])}"
            f"&l={requests.utils.quote(query['l'])}"
            f"&sort=date"
            f"&fromage=1"  # Posted within last 1 day (freshest RSS allows)
        )
        urls.append({"url": url, "source": "Indeed", "country": query["l"]})
    return urls


def build_greenhouse_urls():
    """
    Greenhouse is a hiring platform used by many startups and tech companies.
    Their API is public — no key needed. Returns very fresh listings.
    These are companies in your target sectors known to hire new grads + OPT friendly.
    """
    companies = [
        # Semiconductor & Hardware
        "nvidia", "amd", "qualcomm", "applied-materials", "lam-research",
        "kla", "marvell", "lattice-semiconductor", "monolithic-power-systems",
        # Industrial Automation & Robotics
        "boston-dynamics", "automation-anywhere", "zebra-technologies",
        "cognex", "national-instruments", "rockwell-automation",
        # AI & Software (product engineering focus)
        "openai", "anthropic", "scale-ai", "cohere", "mistral",
        "databricks", "snowflake", "datadog", "figma", "notion",
        "linear", "vercel", "stripe", "airbnb", "doordash",
        # Mid-size startups with rotational programs
        "klaviyo", "hubspot", "asana", "okta", "cloudflare", "hashicorp",
    ]
    urls = []
    for company in companies:
        urls.append({
            "url": f"https://boards-api.greenhouse.io/v1/boards/{company}/jobs",
            "source": "Greenhouse",
            "company": company,
        })
    return urls


# def build_wellfound_urls(queries):
#     """Wellfound (AngelList) — great for startups, has public RSS"""
#     role_map = {
#         "product engineer": "product-engineer",
#         "test engineer": "test-engineer",
#         "applied ai": "ai-engineer",
#         "validation engineer": "validation-engineer",
#     }
#     urls = []
#     for role_key, role_slug in role_map.items():
#         urls.append({
#             "url": f"https://wellfound.com/role/r/{role_slug}?utm_source=rss",
#             "source": "Wellfound",
#             "role": role_key,
#         })
#     return urls


# ──────────────────────────────────────────────
# FRESHNESS FILTER
# ──────────────────────────────────────────────

def is_fresh_enough(pub_date_str, max_hours=24):
    """
    Returns True if the job was posted within max_hours.
    RSS feeds refresh every few hours so 24h is the practical minimum.
    We keep today's jobs and skip anything older.
    """
    if not pub_date_str:
        return True  # If no date, keep it (better safe than miss)
    try:
        pub_dt = parsedate_to_datetime(pub_date_str)
        now = datetime.now(timezone.utc)
        age = now - pub_dt
        return age <= timedelta(hours=max_hours)
    except Exception:
        return True  # If unparseable, keep it


# ──────────────────────────────────────────────
# FETCH FROM INDEED RSS
# ──────────────────────────────────────────────

def fetch_from_indeed(queries):
    jobs = []
    headers = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"}
    urls = build_indeed_urls(queries)

    for entry in urls:
        try:
            response = requests.get(entry["url"], headers=headers, timeout=10)
            if response.status_code != 200:
                print(f"  ⚠️  Indeed returned {response.status_code} for {entry['country']}")
                continue

            root = ET.fromstring(response.content)
            items = root.findall(".//item")
            fresh_count = 0

            for item in items:
                pub_date = item.findtext("pubDate", "")
                if not is_fresh_enough(pub_date, max_hours=24):
                    continue  # Skip old listings

                def clean(text):
                    text = text.replace('\xa0', ' ')
                    return text.encode("ascii", "ignore").decode("ascii").strip()

                jobs.append({
                    "title":       clean(item.findtext("title", "")),
                    "link":        item.findtext("link", "").strip(),
                    "description": clean(item.findtext("description", ""))[:800],
                    "date":        pub_date,
                    "source":      "Indeed",
                    "location":    entry["country"],
                })
                fresh_count += 1

            print(f"  ✅ Indeed [{entry['country']}] → {fresh_count} fresh jobs")
            time.sleep(1)  # Be polite to Indeed's servers

        except Exception as e:
            print(f"  ❌ Indeed error for {entry['country']}: {e}")

    return jobs


# ──────────────────────────────────────────────
# FETCH FROM GREENHOUSE API
# ──────────────────────────────────────────────

def fetch_from_greenhouse():
    jobs = []
    urls = build_greenhouse_urls()

    # Keywords to filter relevant roles from Greenhouse
    relevant_keywords = [
        "product engineer", "test engineer", "validation engineer",
        "applied ai", "ai engineer", "machine learning engineer",
        "rotational", "graduate program", "new grad", "entry level",
        "automation engineer", "hardware engineer", "semiconductor",
    ]

    for entry in urls:
        try:
            response = requests.get(entry["url"], timeout=10)
            if response.status_code != 200:
                continue

            data = response.json()
            all_jobs = data.get("jobs", [])

            for job in all_jobs:
                title = job.get("title", "").lower()

                # Only keep jobs with relevant titles
                if not any(kw in title for kw in relevant_keywords):
                    continue

                location = ""
                if job.get("location"):
                    location = job["location"].get("name", "")

                jobs.append({
                    "title":       job.get("title", ""),
                    "link":        job.get("absolute_url", ""),
                    "description": f"Location: {location}. " + str(job.get("metadata", ""))[:400],
                    "date":        job.get("updated_at", ""),
                    "source":      "Greenhouse",
                    "location":    location,
                    "company":     entry["company"],
                })

            time.sleep(0.5)

        except Exception as e:
            # Many companies won't have a Greenhouse board — that's fine, skip silently
            pass

    print(f"  ✅ Greenhouse → {len(jobs)} relevant jobs found across all companies")
    return jobs


# ──────────────────────────────────────────────
# FETCH FROM GLASSDOOR RSS
# ──────────────────────────────────────────────

def fetch_from_glassdoor(queries):
    jobs = []
    headers = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"}

    # Glassdoor RSS — limited but free
    role_samples = [
        "product+engineer", "test+engineer",
        "validation+engineer", "applied+ai+engineer"
    ]

    for role in role_samples:
        url = f"https://www.glassdoor.com/Job/jobs.htm?suggestCount=0&suggestChosen=false&clickSource=searchBtn&typedKeyword={role}&sc.keyword={role}&locT=N&locId=1&jobType=&context=Jobs&rsRedir=false"
        # Note: Glassdoor doesn't have a true RSS anymore — we use their job search page
        # and fall back to Indeed for Glassdoor-listed jobs (Indeed aggregates them)
        pass

    print(f"  ℹ️  Glassdoor: Using Indeed as aggregator (Glassdoor RSS deprecated in 2023)")
    return jobs  # Indeed already captures Glassdoor-listed jobs


# ──────────────────────────────────────────────
# COUNTRY-SPECIFIC BOARDS
# ──────────────────────────────────────────────

def fetch_country_specific():
    """
    Extra job boards for India, Singapore, Ireland
    that go beyond Indeed's international coverage
    """
    jobs = []
    headers = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"}

    extra_feeds = [
        # Ireland — IrishJobs.ie RSS
        {
            "url": "https://www.irishjobs.ie/TalentAlert/TalentAlert.aspx?SectorID=18&Keyword=engineer",
            "source": "IrishJobs",
            "location": "Dublin, Ireland"
        },
        # Singapore — JobsDB RSS
        {
            "url": "https://sg.jobsdb.com/j?q=product+engineer&l=singapore",
            "source": "JobsDB",
            "location": "Singapore"
        },
        # India — Naukri doesn't have public RSS but Indeed India covers it well
        # Greenhouse also has India offices for many companies above
    ]

    for feed in extra_feeds:
        try:
            response = requests.get(feed["url"], headers=headers, timeout=10)
            if response.status_code == 200:
                root = ET.fromstring(response.content)
                items = root.findall(".//item")
                for item in items:
                    pub_date = item.findtext("pubDate", "")
                    if not is_fresh_enough(pub_date, max_hours=24):
                        continue
                    jobs.append({
                        "title":       item.findtext("title", "").strip(),
                        "link":        item.findtext("link", "").strip(),
                        "description": item.findtext("description", "").strip()[:800],
                        "date":        pub_date,
                        "source":      feed["source"],
                        "location":    feed["location"],
                    })
                print(f"  ✅ {feed['source']} [{feed['location']}] → {len(jobs)} jobs")
        except Exception as e:
            print(f"  ⚠️  {feed['source']} unavailable: {e}")

    return jobs


# ──────────────────────────────────────────────
# MASTER FETCH FUNCTION — calls all sources
# ──────────────────────────────────────────────

def fetch_all_jobs():
    print("\n📡 Fetching jobs from all sources...\n")
    all_jobs = []

    # Build queries only for active countries
    active_queries = []
    for country, is_active in ACTIVE_COUNTRIES.items():
        if is_active:
            for location in SEARCH_LOCATIONS[country]:
                for role in ROLE_KEYWORDS:
                    active_queries.append({"q": role, "l": location})

    # Fetch from each source
    all_jobs += fetch_from_indeed(active_queries)
    all_jobs += fetch_from_greenhouse()
    all_jobs += fetch_country_specific()
    # all_jobs += fetch_from_wellfound()

    # Deduplicate by link
    seen = set()
    unique_jobs = []
    for job in all_jobs:
        key = job.get("link", job.get("title", ""))
        if key and key not in seen:
            seen.add(key)
            unique_jobs.append(job)

    print(f"\n📊 Total unique fresh jobs fetched: {len(unique_jobs)}")
    return unique_jobs

# ──────────────────────────────────────────────
# STEP 2: Filter and enrich jobs using Gemini AI
# ──────────────────────────────────────────────

# Resume summaries for each role — Gemini uses these to score match
# Paste a short summary of each resume here (3-5 bullet points each)
# ──────────────────────────────────────────────
# RESUME LOADER — reads actual .docx files
# ──────────────────────────────────────────────
# Map role categories to resume filenames in the /resumes folder
RESUME_FILES = {
    "product_engineer":         "resumes/01.Kaavya_Sri_Resume_2026.docx",
    "test_validation_engineer": "resumes/01.Kaavya_Sri_Resume_2026_Test_Validation.docx",
    "applied_ai_engineer":      "resumes/01.Kaavya_Sri_Resume_2026_AI_HPC.docx",
}

def load_resume(role_category):
    """
    Reads the .docx resume file for the given role category.
    Returns extracted text, or a fallback message if file not found.
    """
    filepath = RESUME_FILES.get(role_category)

    if not filepath:
        return "Resume not available for this role category."

    if not os.path.exists(filepath):
        print(f"  ⚠️  Resume file not found: {filepath}")
        return f"Resume file missing: {filepath} — please add it to the repo."

    try:
        doc = Document(filepath)
        full_text = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                full_text.append(text)

        # Also extract text from tables (some resumes use table layouts)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        full_text.append(cell_text)

        extracted = "\n".join(full_text)

        # Trim to ~2000 chars to stay within Gemini token limits
        # Keeps the most important top section of your resume
        if len(extracted) > 1000:
            extracted = extracted[:1000] + "\n... [resume trimmed for brevity]"

        print(f"  ✅ Loaded resume: {filepath} ({len(extracted)} chars)")
        return extracted

    except Exception as e:
        print(f"  ❌ Error reading resume {filepath}: {e}")
        return f"Could not read resume file: {e}"


# Pre-load all resumes once at startup (not on every job iteration)
print("\n📄 Loading resume files...")
LOADED_RESUMES = {
    role: load_resume(role)
    for role in RESUME_FILES.keys()
}
print("✅ Resumes loaded\n")


def detect_role_category(job_title, job_description):
    """Map a job listing to one of the three target role categories"""
    title_lower = job_title.lower()
    desc_lower = job_description.lower()
    combined = title_lower + " " + desc_lower

    if any(kw in combined for kw in [
        "test engineer", "validation engineer", "test & validation",
        "test and validation", "hw test", "hardware test", "software test",
        "quality engineer", "reliability engineer", "verification"
    ]):
        return "test_validation_engineer"

    elif any(kw in combined for kw in [
        "applied ai", "ai engineer", "machine learning engineer",
        "ml engineer", "ai applications", "deep learning engineer",
        "computer vision engineer", "nlp engineer"
    ]):
        return "applied_ai_engineer"

    else:
        return "product_engineer"  # default category


def detect_country_context(location_str):
    """Return country context so Gemini applies the right visa filters"""
    loc = location_str.lower()
    if any(x in loc for x in ["india", "hyderabad", "bangalore", "bengaluru"]):
        return "india"
    elif "singapore" in loc:
        return "singapore"
    elif any(x in loc for x in ["ireland", "dublin"]):
        return "ireland"
    else:
        return "usa"  # default


def build_claude_prompt(jobs_batch, country_context):
    """
    Build the filtering + enrichment prompt for gemini.
    Different visa/authorization instructions per country.
    """

    # ── Visa filter instructions vary by country ──────────────
    if country_context == "usa":
        visa_instructions = """
## Visa / Work Authorization Filters (USA ONLY — apply strictly):
The candidate is on OPT (Optional Practical Training) / STEM OPT.
This means:
- They ARE authorized to work in the USA during OPT period (no employer sponsorship needed YET)
- They DO NOT need H1B sponsorship right now
- They ARE NOT a US Citizen or Green Card holder
- They ARE NOT currently enrolled as a student (already graduated)

### HARD SKIP — Remove any job that says ANY of the following:
- "US Citizen only" or "US Citizen required"
- "Green Card required" or "Permanent Resident required"
- "Security clearance required" or "Must hold active clearance"
- "Currently pursuing a degree" or "Must be enrolled in university"
- "Active student enrollment required"
- "No international students"

### FLAG as ⚠️ VERIFY (do NOT auto-skip — candidate decides):
- "No sponsorship available" or "We do not sponsor visas"
  → Reason: This often means no H1B TRANSFER, not necessarily blocking OPT workers
  → Gemini should note: "Says no sponsorship — may still accept OPT. Verify directly."

### SPONSORSHIP LIKELIHOOD BADGE (for USA jobs only):
- ✅ HIGH — Company explicitly says "visa sponsorship available" OR is on known H1B sponsor list
- ⚠️ VERIFY — No mention of sponsorship either way, or says "no sponsorship" (OPT may still be fine)
- ❌ SKIP — Explicitly says "US Citizen only" or "Green Card required" (already filtered out above)
"""
    elif country_context == "india":
        visa_instructions = """
## Work Authorization (INDIA):
- No visa filtering needed — candidate can work freely in India
- Focus filtering on: role relevance, experience level (new grad / 0-2 years), company type
- Flag if job requires 3+ years experience (skip those)
- Prioritize: MNCs with India offices, semiconductor companies, industrial automation firms, startups
"""
    elif country_context == "singapore":
        visa_instructions = """
## Work Authorization (SINGAPORE):
- Candidate would need an Employment Pass (EP) or S Pass for Singapore
- Flag jobs at companies known to hire international candidates
- Large MNCs and tech companies in Singapore commonly sponsor EP
- Skip jobs that say "Singapore Citizens / PRs only"
- Note EP requirement in each listing
"""
    elif country_context == "ireland":
        visa_instructions = """
## Work Authorization (IRELAND / EU):
- Candidate would need a Critical Skills Employment Permit or General Employment Permit
- Engineering roles often qualify for Critical Skills permit
- Large tech companies in Dublin (Google, Meta, LinkedIn, Amazon, Microsoft) regularly sponsor
- Skip if "EU citizens only" or "Right to work in Ireland required without sponsorship"
- Note permit requirement in each listing
"""

    # ── Build jobs text block ──────────────────────────────────
    jobs_text = ""
    for i, job in enumerate(jobs_batch, 1):
        role_cat = detect_role_category(job["title"], job["description"])
        resume = LOADED_RESUMES.get(role_cat, "Resume not available")
        jobs_text += f"""
---
Job #{i}
Title: {job['title']}
Company: {job.get('company', 'See link')}
Location: {job.get('location', 'Not specified')}
Source: {job['source']}
Posted: {job.get('date', 'Unknown')}
Link: {job['link']}
Description: {job['description']}
Detected Role Category: {role_cat.replace('_', ' ').title()}
Candidate Resume for this Role:
{resume}
---
"""

    prompt = f"""
You are an expert job search assistant helping a Masters graduate find engineering roles.
Today's date: {datetime.now().strftime("%B %d, %Y")}

## Candidate Profile:
{YOUR_PROFILE}

{visa_instructions}

## Role Matching Instructions:
The candidate is targeting THREE role types. Match each job to one:
1. **Product Engineer** — building and owning product features, full-stack, platform
2. **Test & Validation Engineer** — hardware/software testing, validation, QA, reliability
3. **Applied AI Engineer** — ML/AI applications in semiconductor or industrial automation

Also flag if a job is:
- 🎓 **Graduate Program** — structured new grad rotational or graduate hire program
- 🔄 **Rotational Program** — multi-team rotation program for new grads

## Company Priority Flags:
Give extra priority to these sectors:
- Semiconductor companies (NVIDIA, AMD, Qualcomm, TI, Applied Materials, Lam Research, KLA, Marvell, etc.)
- Industrial Automation (Siemens, ABB, Rockwell Automation, Honeywell, Cognex, Emerson, Fanuc, etc.)
- Startups and mid-size tech companies (not just FAANG)

## Your Tasks:

### STEP 1 — HARD FILTER
Remove jobs that match the hard skip criteria above.
If fewer than 5 jobs remain after filtering, keep ⚠️ VERIFY ones too.

### STEP 2 — SCORE each remaining job (0–100) based on:
- Role match to one of the 3 target roles (40 pts)
- Resume match — how well candidate's background fits (30 pts)
- Company type match — semiconductor/automation/startup preferred (20 pts)
- Location match to preferred cities (10 pts)

### STEP 3 — ENRICH each job that scores 40 or above with:
1. **Job Title** + **Role Category** (Product / Test & Validation / Applied AI)
2. **Company Name** + sector (Semiconductor / Automation / Startup / Tech)
3. **Team or Department** hiring (if mentioned in description)
4. **Top 3 Requirements** from the listing
5. **Match Score** out of 100 with a 1-sentence reason
6. **Resume Fit** — 1 sentence on how candidate's background fits
7. **Company Overview** — what they do, size, notable products (2-3 sentences)
8. **Work Authorization Note** — based on country context above
9. **Recruiter / Contact Email** — if found in listing (else write "Not listed")
10. **Application Link**
11. **Program Flag** — Graduate Program 🎓 / Rotational 🔄 / Standard Role (pick one)

### STEP 4 — PICK TOP 10 only
After scoring and filtering, keep ONLY the top 10 highest scoring jobs for this country.
If fewer than 10 pass the filters, include all that passed.
Quality over quantity — do not pad with weak matches.

### STEP 5 — FIND HIRING CONTACT
For each job, try to extract from the listing text:
- Hiring manager name and email
- Recruiter name and email
- General HR/careers email (e.g. careers@company.com, jobs@company.com)
If none found in the listing, write the most likely careers email based on company name
e.g. for Siemens write "careers@siemens.com (inferred)" so candidate can verify.

### STEP 6 — BUILD TWO OUTPUTS

#### OUTPUT A — Quick Snapshot Table (goes FIRST in email)
A compact HTML table summarizing all matched jobs:

<table style="width:100%; border-collapse:collapse; font-size:13px; margin-bottom:32px;">
  <thead>
    <tr style="background:#1a1a2e; color:white;">
      <th style="padding:10px; text-align:left;">#</th>
      <th style="padding:10px; text-align:left;">Job Title</th>
      <th style="padding:10px; text-align:left;">Company</th>
      <th style="padding:10px; text-align:left;">Role Type</th>
      <th style="padding:10px; text-align:left;">Location</th>
      <th style="padding:10px; text-align:left;">Score</th>
      <th style="padding:10px; text-align:left;">Work Auth</th>
      <th style="padding:10px; text-align:left;">Posted</th>
      <th style="padding:10px; text-align:left;">Apply</th>
    </tr>
  </thead>
  <tbody>
    [one row per job — alternate row background #f9f9f9 and #ffffff]
    [Score cell: color green if 75+, orange if 50-74, red if below 50]
    [Posted cell: show 🔥 Fresh if posted within 24 hours]
    [Apply cell: small "Apply →" hyperlink]
  </tbody>
</table>

#### OUTPUT B — Detailed Job Cards (goes AFTER the table)
One card per job using exactly this template:

<div style="border:1px solid #e0e0e0; border-radius:10px; padding:20px; margin-bottom:28px; font-family:Arial,sans-serif; background:#ffffff;">

  <!-- Header Row -->
  <div style="display:flex; justify-content:space-between; align-items:flex-start; flex-wrap:wrap; gap:8px; margin-bottom:12px;">
    <div>
      <h2 style="color:#1a1a2e; margin:0 0 4px 0; font-size:18px;">[Job Title] [🔥 if posted within 24h]</h2>
      <h3 style="color:#4a4a8a; margin:0; font-size:14px; font-weight:normal;">[Company] · [Sector: Semiconductor / Automation / Startup / Tech]</h3>
    </div>
    <div style="text-align:right;">
      <span style="background:#f0f0ff; color:#4a4a8a; padding:4px 12px; border-radius:20px; font-size:13px; font-weight:bold;">[Role Category]</span>
      <br>
      <span style="font-size:12px; color:#999; margin-top:4px; display:block;">Posted: [date]</span>
    </div>
  </div>

  <hr style="border:none; border-top:1px solid #f0f0f0; margin:0 0 14px 0;">

  <!-- Match Score Banner -->
  <div style="background:#f8f8ff; border-left:4px solid #4a4a8a; padding:10px 14px; border-radius:0 6px 6px 0; margin-bottom:16px;">
    <strong style="font-size:15px; color:#1a1a2e;">Match Score: [X]/100</strong>
    <p style="margin:4px 0 0 0; font-size:13px; color:#555;">[2 sentence breakdown — why this score: role fit + resume fit + company type + location]</p>
  </div>

  <!-- Job Details Table -->
  <table style="width:100%; font-size:14px; border-collapse:collapse; margin-bottom:16px;">
    <tr style="border-bottom:1px solid #f5f5f5;">
      <td style="padding:8px 12px 8px 0; color:#888; width:170px; vertical-align:top;">📍 Location</td>
      <td style="padding:8px 0; vertical-align:top;">[Location — Remote / City]</td>
    </tr>
    <tr style="border-bottom:1px solid #f5f5f5;">
      <td style="padding:8px 12px 8px 0; color:#888; vertical-align:top;">🏢 Team / Dept</td>
      <td style="padding:8px 0; vertical-align:top;">[Team or Department name — e.g. "Platform Engineering", "Semiconductor Test Div."]</td>
    </tr>
    <tr style="border-bottom:1px solid #f5f5f5;">
      <td style="padding:8px 12px 8px 0; color:#888; vertical-align:top;">📋 Key Requirements</td>
      <td style="padding:8px 0; vertical-align:top;">[req1] · [req2] · [req3] · [req4]</td>
    </tr>
    <tr style="border-bottom:1px solid #f5f5f5;">
      <td style="padding:8px 12px 8px 0; color:#888; vertical-align:top;">📄 JD Summary</td>
      <td style="padding:8px 0; vertical-align:top;">[3-4 sentence summary of the full job description — what the role does day to day, what team you join, what you build or test]</td>
    </tr>
    <tr style="border-bottom:1px solid #f5f5f5;">
      <td style="padding:8px 12px 8px 0; color:#888; vertical-align:top;">💼 Resume Fit</td>
      <td style="padding:8px 0; vertical-align:top;">[2 sentences — specifically how candidate's resume matches this JD. Mention actual skills/projects from resume that align]</td>
    </tr>
    <tr style="border-bottom:1px solid #f5f5f5;">
      <td style="padding:8px 12px 8px 0; color:#888; vertical-align:top;">🏭 About the Company</td>
      <td style="padding:8px 0; vertical-align:top;">[4-5 sentences: what the company does, their main products, company size, notable achievements, why it is a good company for new grads in this field]</td>
    </tr>
    <tr style="border-bottom:1px solid #f5f5f5;">
      <td style="padding:8px 12px 8px 0; color:#888; vertical-align:top;">👥 Hiring Team</td>
      <td style="padding:8px 0; vertical-align:top;">[Team description — what this team works on, their mission, tech stack if mentioned]</td>
    </tr>
    <tr style="border-bottom:1px solid #f5f5f5;">
      <td style="padding:8px 12px 8px 0; color:#888; vertical-align:top;">✉️ Hiring Contact</td>
      <td style="padding:8px 0; vertical-align:top;">[Name + email if found in listing] OR [inferred careers email — clearly marked as "inferred, verify before sending"]</td>
    </tr>
    <tr>
      <td style="padding:8px 12px 8px 0; color:#888; vertical-align:top;">🛂 Work Auth</td>
      <td style="padding:8px 0; vertical-align:top;">[Sponsorship badge + plain English note e.g. "✅ HIGH — explicitly mentions visa sponsorship" or "⚠️ VERIFY — no mention found, OPT likely accepted"]</td>
    </tr>
  </table>

  <!-- Program Flag if applicable -->
  [IF graduate or rotational program add this block:]
  <div style="background:#fff8e1; border:1px solid #ffc107; border-radius:6px; padding:10px 14px; margin-bottom:14px; font-size:13px;">
    🎓 <strong>Graduate / Rotational Program</strong> — [1 sentence describing the program structure e.g. "3 rotation, 18 month program across hardware, software and product teams"]
  </div>

  <!-- Apply Button -->
  <a href="[application link]" style="display:inline-block; background:#1a1a2e; color:white; padding:10px 22px; border-radius:6px; text-decoration:none; font-size:14px; margin-right:10px;">Apply Now →</a>
  <a href="[company careers page if known]" style="display:inline-block; background:#f0f0ff; color:#4a4a8a; padding:10px 22px; border-radius:6px; text-decoration:none; font-size:14px;">Company Careers Page →</a>

</div>

Start with this summary line before the cards:
<p style="font-size:15px; color:#444;">Found <strong>[X] matching roles</strong> out of [total] listings reviewed — [Date]. Sorted by work auth compatibility then match score.</p>

If zero jobs match after filtering, return:
<p style="color:#888;">No new matching listings found in this batch. Next run scheduled soon.</p>

## Job Listings to Analyze:
{jobs_text}
"""
    return prompt


def analyze_jobs_with_claude(all_jobs):
    """
    Splits jobs into batches by country context,
    sends each batch to Gemini with the right filters,
    combines results into one email.
    """
    if not all_jobs:
        print("No jobs to analyze.")
        return "<p style='color:#888;'>No new job listings found today.</p>"

    genai.configure(api_key=os.environ["GEMINI_API_KEY"])
    # # --- DIAGNOSTIC: Print available models ---
    # print("\n🔍 Checking available Gemini models for your API key:")
    # try:
    #     for m in genai.list_models():
    #         if 'generateContent' in m.supported_generation_methods:
    #             print(f"  ⭐ {m.name}")
    # except Exception as e:
    #     print(f"  ❌ Could not list models: {e}")
    # ------------------------------------------
    # Update this line
    model = genai.GenerativeModel("gemini-2.0-flash")


    # ── Group jobs by country context ──
    country_batches = {"usa": [], "india": [], "singapore": [], "ireland": []}
    for job in all_jobs:
        ctx = detect_country_context(job.get("location", ""))
        country_batches[ctx].append(job)

    all_html_sections = []

    country_labels = {
        "usa":       "🇺🇸 United States",
        "india":     "🇮🇳 India",
        "singapore": "🇸🇬 Singapore",
        "ireland":   "🇮🇪 Ireland",
    }

    for country, jobs in country_batches.items():
        if not jobs:
            continue

        # Skip if country not active
        if not ACTIVE_COUNTRIES.get(country, False):
            continue

        print(f"\n🤖 Analyzing {len(jobs)} jobs for {country_labels[country]}...")

        # Batch into chunks of 30 (token limit safety)
        batch_size = 10
        country_html_parts = []

        for i in range(0, len(jobs), batch_size):
            batch = jobs[i:i + batch_size]
            batch_num = (i // batch_size) + 1
            total_batches = (len(jobs) + batch_size - 1) // batch_size
            print(f"  📦 Batch {batch_num}/{total_batches} — {len(batch)} jobs")

            try:
                prompt = build_claude_prompt(batch, country)
                response = model.generate_content(prompt)
                country_html_parts.append(response.text)
                time.sleep(10)  # Gemini free tier — slightly longer pause 

            except Exception as e:
                print(f"  ❌ Gemini error on batch {batch_num}: {e}")
                country_html_parts.append(
                    f"<p style='color:red;'>Error analyzing batch {batch_num}: {e}</p>"
                )

        # Wrap country section with header
        country_section = f"""
<div style="margin-bottom:40px;">
  <div style="background:#1a1a2e; color:white; padding:14px 20px; border-radius:8px; margin-bottom:20px;">
    <h2 style="margin:0; font-size:18px;">{country_labels[country]}</h2>
    <p style="margin:4px 0 0 0; opacity:0.7; font-size:13px;">{len(jobs)} listings reviewed</p>
  </div>
  {''.join(country_html_parts)}
</div>
"""
        all_html_sections.append(country_section)
        print(f"  ✅ {country_labels[country]} analysis complete")

    if not all_html_sections:
        return "<p style='color:#888;'>No matching jobs found across any region today.</p>"

    print("\n✅ All Gemini analysis complete")
    return "\n".join(all_html_sections)

# ──────────────────────────────────────────────
# STEP 3: Send the email digest
# ──────────────────────────────────────────────

def send_email(html_content):
    sender_email    = os.environ["EMAIL_ADDRESS"]
    sender_password = os.environ["EMAIL_APP_PASSWORD"]
    recipient_email = os.environ["EMAIL_TO"]

    today = datetime.now().strftime("%B %d, %Y")
    run_time = datetime.now().strftime("%Y-%m-%d %H:%M UTC")

    msg = MIMEMultipart("alternative")
    msg["Subject"] = f"🚀 Engineering Job Digest — Top 10 Per Region | {today}"
    msg["From"]    = sender_email
    msg["To"]      = recipient_email

    plain_text = (
        "Your daily engineering job digest is ready. "
        "Please open in an HTML email client to view formatted results."
    )

    full_html = f"""
    <html>
    <body style="font-family:Arial,sans-serif; max-width:760px; margin:auto; padding:20px; background:#f4f4f8;">

      <!-- ── EMAIL HEADER ── -->
      <div style="background:linear-gradient(135deg,#1a1a2e,#4a4a8a);
                  color:white; padding:28px 24px; border-radius:12px; margin-bottom:24px;">
        <h1 style="margin:0 0 6px 0; font-size:22px;">🚀 Daily Engineering Job Digest</h1>
        <p style="margin:0; opacity:0.8; font-size:14px;">
          Product Engineer · Test &amp; Validation · Applied AI
          &nbsp;|&nbsp; OPT/STEM OPT Friendly &nbsp;|&nbsp; {today}
        </p>
      </div>

      <!-- ── HOW TO READ THIS EMAIL ── -->
      <div style="background:#e8f4fd; border:1px solid #90caf9; border-radius:8px;
                  padding:14px 18px; margin-bottom:24px; font-size:13px; color:#1a1a2e;">
        <strong>📖 How to use this digest:</strong><br>
        1. Scan the <strong>Quick Snapshot Table</strong> at the top of each country section<br>
        2. Click into cards with <strong>score 75+</strong> first — those are your best matches<br>
        3. Jobs marked <strong>🔥 Fresh</strong> were posted in the last 24 hours — apply same day<br>
        4. Check <strong>Work Auth</strong> row before applying — ✅ HIGH = confirmed friendly, ⚠️ VERIFY = confirm with recruiter<br>
        5. Use the <strong>Hiring Contact</strong> email to reach out directly after applying
      </div>

      <!-- ── OPT REMINDER ── -->
      <div style="background:#fff8e1; border:1px solid #ffc107; border-radius:8px;
                  padding:12px 18px; margin-bottom:28px; font-size:13px;">
        <strong>⚠️ OPT Reminder:</strong>
        Always verify work authorization directly with the recruiter.
        "No sponsorship" often means no H1B transfer — OPT workers are frequently still accepted.
        Check company H1B history at
        <a href="https://www.myvisajobs.com" style="color:#4a4a8a;">myvisajobs.com</a>
        before applying.
      </div>

      <!-- ── MAIN CONTENT (Gemini's output per country) ── -->
      {html_content}

      <!-- ── APPLICATION CHECKLIST FOOTER ── -->
      <div style="background:#1a1a2e; color:white; border-radius:10px;
                  padding:20px 24px; margin-top:32px;">
        <h3 style="margin:0 0 8px 0; font-size:16px;">✅ Application Checklist</h3>
        <p style="margin:0 0 12px 0; font-size:13px; opacity:0.8;">
          Before applying to any role, make sure you have:
        </p>
        <ul style="margin:0; padding-left:20px; font-size:13px; opacity:0.9; line-height:1.9;">
          <li>Tailored your resume to match the specific role (Product / Test / Applied AI)</li>
          <li>Verified the work authorization requirement with the recruiter</li>
          <li>Checked the company's H1B sponsorship history on myvisajobs.com</li>
          <li>Written a short outreach email to the hiring contact if email was found</li>
          <li>Noted your OPT end date and STEM OPT eligibility in your cover note if asked</li>
        </ul>
      </div>

      <!-- ── EMAIL FOOTER ── -->
      <div style="text-align:center; color:#aaa; font-size:12px;
                  margin-top:24px; padding-top:16px; border-top:1px solid #ddd;">
        <p style="margin:0;">
          Generated automatically using Gemini AI + GitHub Actions<br>
          Run time: {run_time} &nbsp;·&nbsp;
          Regions: USA 🇺🇸 · India 🇮🇳 · Singapore 🇸🇬 · Ireland 🇮🇪
        </p>
      </div>

    </body>
    </html>
    """

    msg.attach(MIMEText(plain_text, "plain", "utf-8"))
    msg.attach(MIMEText(full_html, "html", "utf-8"))

    try:
        with smtplib.SMTP_SSL("smtp.gmail.com", 465) as server:
            server.login(sender_email, sender_password)
            server.sendmail(sender_email, recipient_email, msg.as_bytes())
        print(f"✅ Email sent to {recipient_email}")
    except Exception as e:
        print(f"❌ Failed to send email: {e}")
        raise
# ```

# ---

# ## What the Email Looks Like End to End
# ```
# 📧 Subject: 🚀 Engineering Job Digest — Top 10 Per Region | March 1, 2025
# ─────────────────────────────────────────────────
#   HEADER BANNER — Role types + date

#   HOW TO READ THIS — quick guide

#   OPT REMINDER — myvisajobs.com tip

#   ── 🇺🇸 United States ──────────────────────────
#     QUICK SNAPSHOT TABLE (all 10 in one view)
#     #  Title  Company  Role  Location  Score  Auth  Posted  Apply

#     DETAILED CARD × 10
#     each card has:
#       Match Score Banner (score + reason)
#       JD Summary
#       Resume Fit
#       Company Deep Dive (4-5 sentences)
#       Hiring Team description
#       Hiring Contact email
#       Work Auth badge
#       Apply Now + Careers Page buttons

#   ── 🇮🇳 India ──────────────────────────────────
#     [same structure × 10]

#   ── 🇸🇬 Singapore ──────────────────────────────
#     [same structure × 10]

#   ── 🇮🇪 Ireland ────────────────────────────────
#     [same structure × 10]

#   APPLICATION CHECKLIST FOOTER
# ─────────────────────────────────────────────────

# ──────────────────────────────────────────────
# MAIN — Ties the full pipeline together
# ──────────────────────────────────────────────

def run_pipeline():
    """
    Full job search pipeline:
    1. Load resumes
    2. Fetch jobs from all sources + countries
    3. Analyze + filter + score with Gemini
    4. Send email digest
    5. Log summary report
    """

    start_time = datetime.now()

    print("\n" + "="*60)
    print("  🚀 Engineering Job Search Bot")
    print(f"  📅 {start_time.strftime('%A, %B %d, %Y — %H:%M UTC')}")
    print("="*60)

    # ── Active countries summary ──────────────────────────────
    active = [c.upper() for c, v in ACTIVE_COUNTRIES.items() if v]
    inactive = [c.upper() for c, v in ACTIVE_COUNTRIES.items() if not v]
    print(f"\n🌍 Active regions   : {', '.join(active) if active else 'None'}")
    print(f"💤 Skipped regions  : {', '.join(inactive) if inactive else 'None'}")

    # ── Resume check ─────────────────────────────────────────
    print("\n📄 Resume status:")
    resumes_ok = True
    for role, text in LOADED_RESUMES.items():
        status = "✅" if "missing" not in text.lower() and "error" not in text.lower() else "❌"
        print(f"  {status} {role.replace('_', ' ').title()}")
        if status == "❌":
            resumes_ok = False

    if not resumes_ok:
        print("\n  ⚠️  One or more resume files missing.")
        print("  Script will continue but match scoring may be less accurate.")
        print("  Add .docx files to /resumes folder and push to repo.\n")

    # ── Check required environment variables ─────────────────
    print("\n🔐 Checking environment secrets...")
    required_secrets = [
        "GEMINI_API_KEY",
        "EMAIL_ADDRESS",
        "EMAIL_APP_PASSWORD",
        "EMAIL_TO",
    ]
    missing_secrets = []
    for secret in required_secrets:
        if not os.environ.get(secret):
            missing_secrets.append(secret)
            print(f"  ❌ Missing: {secret}")
        else:
            print(f"  ✅ Found  : {secret}")

    if missing_secrets:
        print(f"\n❌ FATAL: Missing required secrets: {', '.join(missing_secrets)}")
        print("   Go to GitHub → Settings → Secrets and Variables → Actions")
        print("   Add the missing secrets and re-run.\n")
        sys.exit(1)

    # ─────────────────────────────────────────────────────────
    # STEP 1 — Fetch all jobs
    # ─────────────────────────────────────────────────────────
    print("\n" + "-"*60)
    print("STEP 1 — Fetching jobs from all sources")
    print("-"*60)

    try:
        all_jobs = fetch_all_jobs()
    except Exception as e:
        print(f"\n❌ Fatal error during job fetching: {e}")
        send_error_email("Job Fetching Failed", str(e))
        sys.exit(1)

    if not all_jobs:
        print("\n⚠️  No jobs fetched from any source.")
        send_error_email(
            "No Jobs Found",
            "The job search bot ran but found no listings from any source. "
            "Indeed or other feeds may be temporarily unavailable."
        )
        sys.exit(0)

    # ── Per-country breakdown ─────────────────────────────────
    print("\n📊 Jobs fetched per region:")
    country_counts = {}
    for job in all_jobs:
        ctx = detect_country_context(job.get("location", ""))
        country_counts[ctx] = country_counts.get(ctx, 0) + 1
    for country, count in country_counts.items():
        print(f"  🌍 {country.upper():<12} → {count} jobs")
    print(f"  {'TOTAL':<14} → {len(all_jobs)} jobs")

    # ─────────────────────────────────────────────────────────
    # STEP 2 — Analyze + filter + score with Gemini
    # ─────────────────────────────────────────────────────────
    print("\n" + "-"*60)
    print("STEP 2 — Analyzing jobs with Gemini AI")
    print("-"*60)

    try:
        analysis_html = analyze_jobs_with_claude(all_jobs)
    except Exception as e:
        print(f"\n❌ Fatal error during Gemini analysis: {e}")
        send_error_email("Gemini Analysis Failed", str(e))
        sys.exit(1)

    # Check if Gemini returned empty / no matches
    if not analysis_html or "No matching jobs found" in analysis_html:
        print("\n⚠️  Gemini found no matching jobs after filtering.")
        send_no_matches_email()
        sys.exit(0)

    # ─────────────────────────────────────────────────────────
    # STEP 3 — Send email digest
    # ─────────────────────────────────────────────────────────
    print("\n" + "-"*60)
    print("STEP 3 — Sending email digest")
    print("-"*60)

    try:
        send_email(analysis_html)
    except Exception as e:
        print(f"\n❌ Fatal error sending email: {e}")
        sys.exit(1)

    # ─────────────────────────────────────────────────────────
    # DONE — Final summary log
    # ─────────────────────────────────────────────────────────
    end_time   = datetime.now()
    total_time = (end_time - start_time).seconds

    print("\n" + "="*60)
    print("  ✅ Pipeline Complete!")
    print("="*60)
    print(f"  📬 Email sent to   : {os.environ.get('EMAIL_TO')}")
    print(f"  📦 Jobs fetched    : {len(all_jobs)}")
    print(f"  🌍 Regions covered : {', '.join(active)}")
    print(f"  ⏱️  Total runtime   : {total_time}s")
    print(f"  🕐 Finished at     : {end_time.strftime('%H:%M:%S UTC')}")
    print("="*60 + "\n")


# ──────────────────────────────────────────────
# ERROR EMAIL HELPERS
# Helper emails so you always know if something breaks
# ──────────────────────────────────────────────

def send_error_email(error_type, error_detail):
    """
    Sends a plain alert email if any step of the pipeline fails.
    So you know the bot broke rather than just getting no email silently.
    """
    try:
        sender_email    = os.environ.get("EMAIL_ADDRESS", "")
        sender_password = os.environ.get("EMAIL_APP_PASSWORD", "")
        recipient_email = os.environ.get("EMAIL_TO", "")

        if not all([sender_email, sender_password, recipient_email]):
            print("⚠️  Cannot send error email — email secrets not set")
            return

        msg = MIMEMultipart("alternative")
        msg["Subject"] = f"⚠️ Job Bot Error: {error_type}"
        msg["From"]    = sender_email
        msg["To"]      = recipient_email

        html = f"""
        <html>
        <body style="font-family:Arial,sans-serif; max-width:600px; margin:auto; padding:20px;">
          <div style="background:#ff4444; color:white; padding:20px; border-radius:10px; margin-bottom:20px;">
            <h2 style="margin:0;">⚠️ Job Search Bot Error</h2>
            <p style="margin:8px 0 0 0; opacity:0.9;">{error_type}</p>
          </div>
          <div style="background:#fff3f3; border:1px solid #ffcccc; border-radius:8px; padding:16px;">
            <p><strong>Error Detail:</strong></p>
            <pre style="background:#f9f9f9; padding:12px; border-radius:6px; font-size:13px; white-space:pre-wrap;">{error_detail}</pre>
          </div>
          <div style="margin-top:20px; font-size:13px; color:#666;">
            <p>To debug:</p>
            <ol>
              <li>Go to your GitHub repo → <strong>Actions</strong> tab</li>
              <li>Click the failed workflow run</li>
              <li>Expand the <strong>Run job search bot</strong> step for full logs</li>
            </ol>
          </div>
        </body>
        </html>
        """

        msg.attach(MIMEText(html, "html", "utf-8"))

        with smtplib.SMTP_SSL("smtp.gmail.com", 465) as server:
            server.login(sender_email, sender_password)
            server.sendmail(sender_email, recipient_email, msg.as_bytes())

        print(f"⚠️  Error notification sent to {recipient_email}")

    except Exception as e:
        print(f"❌ Could not send error email either: {e}")


def send_no_matches_email():
    """
    Sends a short email when bot ran fine but found zero matching jobs today.
    So silence = actual silence, not a broken bot.
    """
    try:
        sender_email    = os.environ.get("EMAIL_ADDRESS", "")
        sender_password = os.environ.get("EMAIL_APP_PASSWORD", "")
        recipient_email = os.environ.get("EMAIL_TO", "")

        today = datetime.now().strftime("%B %d, %Y")

        msg = MIMEMultipart("alternative")
        msg["Subject"] = f"📭 Job Bot — No New Matches Today | {today}"
        msg["From"]    = sender_email
        msg["To"]      = recipient_email

        html = f"""
        <html>
        <body style="font-family:Arial,sans-serif; max-width:600px; margin:auto; padding:20px;">
          <div style="background:#f0f0ff; border:1px solid #c0c0ff;
                      border-radius:10px; padding:24px; text-align:center;">
            <h2 style="color:#4a4a8a; margin:0 0 8px 0;">📭 No New Matches Today</h2>
            <p style="color:#666; margin:0;">{today}</p>
          </div>
          <div style="margin-top:20px; font-size:14px; color:#444; line-height:1.7;">
            <p>The bot ran successfully across all active regions but found
            no new listings that passed the filters today.</p>
            <p>This is normal — some days are quiet. The bot will run again
            on the next scheduled cycle.</p>
            <p style="margin-top:16px;"><strong>Active regions checked:</strong><br>
            🇺🇸 USA &nbsp;·&nbsp; 🇮🇳 India &nbsp;·&nbsp;
            🇸🇬 Singapore &nbsp;·&nbsp; 🇮🇪 Ireland</p>
          </div>
        </body>
        </html>
        """

        msg.attach(MIMEText(html, "html", "utf-8"))

        with smtplib.SMTP_SSL("smtp.gmail.com", 465) as server:
            server.login(sender_email, sender_password)
            server.sendmail(sender_email, recipient_email, msg.as_bytes())

        print(f"📭 No-matches notification sent to {recipient_email}")

    except Exception as e:
        print(f"❌ Could not send no-matches email: {e}")


# ──────────────────────────────────────────────
# ENTRY POINT
# ──────────────────────────────────────────────

if __name__ == "__main__":
    run_pipeline()

